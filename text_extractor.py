from typing import List, Optional, Union, Tuple
from pathlib import Path
import fitz  # PyMuPDF
import docx
from loguru import logger
from langchain.schema.document import Document
from PIL import Image
import pytesseract
import re

class TextExtractor:
    """
    Класс для извлечения текста из различных форматов файлов (PDF, DOCX).
    Поддерживает OCR для PDF-файлов, не содержащих текстового слоя.
    """

    def __init__(self):
        logger.info("Инициализирован TextExtractor.")

    def _extract_gdrive_links(self, text: str) -> List[str]:
        """Извлекает ID папок Google Drive из текста с помощью регулярного выражения."""
        # Паттерн для поиска URL вида https://drive.google.com/drive/folders/ID_ПАПКИ
        # и захвата только самого ID
        pattern = r"https://drive\.google\.com/drive/folders/([a-zA-Z0-9_-]+)"
        return re.findall(pattern, text)

    def extract_from_file(self, file_path: Union[str, Path]) -> Tuple[Optional[Document], List[str]]:
        """
        Извлекает текст и ссылки на папки GDrive из файла.

        Args:
            file_path (Union[str, Path]): Путь к файлу.

        Returns:
            Tuple[Optional[Document], List[str]]: Кортеж, где первый элемент - 
                объект Document с текстом, а второй - список ID найденных папок.
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        logger.info(f"Начало извлечения текста из файла: {file_path.name}")

        text = ""
        try:
            if extension == ".pdf":
                text = self._extract_pdf_with_ocr(file_path)
            elif extension == ".docx":
                text = self._extract_docx(file_path)
            else:
                logger.warning(f"Неподдерживаемый формат файла: {extension}. Файл будет пропущен.")
                return None, []

            if not text or text.isspace():
                logger.warning(f"Документ '{file_path.name}' пуст или содержит только пробелы. OCR также не дал результатов.")
                return None, []

            # Извлекаем ссылки из полученного текста
            links = self._extract_gdrive_links(text)
            if links:
                logger.info(f"Найдено {len(links)} ссылок на папки Google Drive в файле '{file_path.name}'.")

            metadata = {"source": file_path.name, "full_path": str(file_path)}
            doc = Document(page_content=text, metadata=metadata)
            logger.success(f"Текст из файла '{file_path.name}' успешно извлечен.")
            return doc, links

        except Exception as e:
            logger.error(f"Не удалось извлечь текст из файла {file_path.name}: {e}")
            return None, []

    def _extract_pdf_with_ocr(self, file_path: Path) -> str:
        """
        Извлекает текст из PDF. Если текстовый слой пуст, применяет OCR.
        """
        text = ""
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
        except Exception as e:
            logger.error(f"Ошибка при первичном чтении PDF '{file_path.name}': {e}")
            return "" # Возвращаем пустую строку, чтобы не падать

        # Если после обычного извлечения текста нет, пробуем OCR
        if not text.strip():
            logger.info(f"Текстовый слой в '{file_path.name}' пуст. Запускаю OCR...")
            try:
                ocr_text = ""
                with fitz.open(file_path) as doc:
                    for i, page in enumerate(doc):
                        logger.debug(f"Обработка страницы {i+1}/{len(doc)} через OCR...")
                        # Уменьшаем разрешение до 200 DPI для снижения потребления памяти
                        pix = page.get_pixmap(dpi=200)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        # Распознаем текст, указывая русский язык
                        ocr_text += pytesseract.image_to_string(img, lang='rus') + "\n"
                
                if ocr_text.strip():
                    logger.success(f"OCR успешно завершен для '{file_path.name}'.")
                    return ocr_text
                else:
                    logger.warning(f"OCR не смог распознать текст в '{file_path.name}'.")
                    return "" # Возвращаем пустую строку, если OCR ничего не нашел
            except Exception as e:
                logger.error(f"Ошибка в процессе OCR для файла '{file_path.name}': {e}")
                return "" # В случае ошибки OCR, возвращаем пустую строку

        return text

    def _extract_docx(self, file_path: Path) -> str:
        """Извлекает текст из DOCX-файла."""
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

# Пример использования
if __name__ == '__main__':
    logger.info("--- Тестирование TextExtractor ---")
    
    # Создаем временную директорию для тестов
    test_dir = Path("./test_files")
    test_dir.mkdir(exist_ok=True)
    
    extractor = TextExtractor()

    # 1. Тест с DOCX
    docx_path = test_dir / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("Это первая строка в документе Word.")
    doc.add_paragraph("А это вторая строка.")
    doc.save(docx_path)
    
    extracted_docx, links = extractor.extract_from_file(docx_path)
    if extracted_docx:
        print("\n--- Результат для DOCX ---")
        print(f"Извлеченный текст:\n{extracted_docx.page_content}")
        print(f"Метаданные: {extracted_docx.metadata}")
        print(f"Найденные ссылки на папки GDrive: {links}")

    # 2. Тест с PDF (создание PDF налету - сложно, лучше проверить с существующим)
    # Предполагается, что у вас есть файл test.pdf. Если нет, этот тест будет пропущен.
    pdf_path = test_dir / "test.pdf"
    if pdf_path.exists():
        extracted_pdf, links = extractor.extract_from_file(pdf_path)
        if extracted_pdf:
            print("\n--- Результат для PDF ---")
            print(f"Извлеченный текст (первые 100 символов): '{extracted_pdf.page_content[:100]}...'")
            print(f"Метаданные: {extracted_pdf.metadata}")
            print(f"Найденные ссылки на папки GDrive: {links}")
    else:
        print(f"\nДля теста с PDF, пожалуйста, поместите файл 'test.pdf' в директорию '{test_dir}'")

    # 3. Тест с неподдерживаемым форматом
    txt_path = test_dir / "test.txt"
    txt_path.write_text("простой текст")
    extracted_txt, links = extractor.extract_from_file(txt_path)
    print("\n--- Результат для TXT (неподдерживаемый) ---")
    print(f"Результат: {extracted_txt}")
    print(f"Найденные ссылки на папки GDrive: {links}")


    # Очистка
    import shutil
    shutil.rmtree(test_dir)
    logger.info(f"\nТестовая директория '{test_dir}' удалена.")
    logger.info("Тестирование завершено.")
